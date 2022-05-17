from codecs import namereplace_errors
from enum import Enum
from pydoc import doc
from docx import Document
import gender_guesser.detector as gender

class Gender(Enum):
	UNKNOWN: 'unknown'
	ANDY: 'andy'
	MALE: 'male'
	FEMALE: 'female'
	MOSTLY_MALE: 'mostly_male'
	MOSTLY_FEMALE: 'mostly_female'




def get_addressing(name):
	g = gender.Detector()
	mf = g.get_gender(name)

	#print(mf)
	if mf == 'unknown':
		# TODO: warning
		return ''
	elif mf == 'andy':
		# TODO: warning
		return ''
	elif mf == 'male':
		return 'heer '
	elif mf == 'female':
		return 'mevrouw '
	elif mf == 'mostly_male':
		# TODO: warning
		pass
	elif mf == 'mostly_female':
		# TODO: warning
		pass
	else:
		# TODO: ERROR
		pass
	return ''



def generate_doc(name_police, date_received, date_of_offense, uur, name):
	document = Document()

	addressing = get_addressing(name)

	# TODO
	addressing = ''

	document.add_paragraph(f'Geachte {addressing}{name_police},')
	# TODO: fix datum wanneer de pv is doorgestuurd/ontvangen
	document.add_paragraph(f'Via deze brief wil ik graag wat meer duidelijkheid scheppen en mijn verweer indienen tegen de gemeentelijke administratieve geldboete die mij onlangs werd toegestuurd. Dit betreft het overtreden van het politieregelement (artikel 769.19).')
		#document.add_paragraph(f'Via deze brief wil ik graag wat meer duidelijkheid scheppen en mijn verweer indienen tegen de gemeentelijke administratieve geldboete die mij op {date_received} werd toegestuurd. Dit betreft het overtreden van het politieregelement (artikel 769.19).')
	document.add_paragraph(f'Allereerst wil ik beginnen met waarom ik op {date_of_offense} om {uur} met mijn voertuig door Parijsstraat 72 te Leuven reed. Sinds mei 2019 ben ik tewerk gesteld bij Mise en Place Leuven, gevestigd Parijsstraat 76, te Leuven. De parking van Mise en Place is gesitueerd op hetzelfde adres. Het is vanzelfsprekend dat ik in de Parijsstraat passeer om tot aan de parking te komen.')
	document.add_paragraph('Omdat ik nog student ben en geen vast werknemer, is mijn nummerplaat niet geregistreerd. Dus elke keer als ik aldaar moet zijn, bel ik een van de verantwoordelijke bij Mise en Place op en zeg dat ik op datum x op de site van Mise en Place moet zijn. Deze geeft vervolgens mijn nummerplaat door aan jullie. Ook nu dacht ik dat ik mijn nummerplaat had doorgegeven, niets bleek minder waar te zijn toen ik jullie proces-verbaal toegestuurd kreeg.')
	document.add_paragraph('Aarzel niet mij telefonisch te contacteren op volgend nummer +324 92/956.496. Of via email, indien u nog verdere informatie nodig heeft.')
	document.add_paragraph('Hoogachtend,\nArno Barré')

	document.add_page_break()

	document.save(f"verweer/verweer_{str(date_of_offense).replace('-', '')}.docx")
	print(f"SAVED: verweer_{str(date_of_offense).replace('-', '')}.docx")


